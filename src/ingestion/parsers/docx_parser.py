from typing import Optional

from src.ingestion.parsers.text_parser import TextParser


class DocxParser:
    """Shim parser for DOCX inputs that delegates extracted text to TextParser."""

    def __init__(self, text_parser: Optional[TextParser] = None) -> None:
        self.text_parser = text_parser or TextParser()

    def parse(self, document) -> dict:
        return self.text_parser.parse(self._extract_text(document))

    def _extract_text(self, document) -> str:
        if hasattr(document, "paragraphs"):
            return "\n\n".join(paragraph.text for paragraph in document.paragraphs)

        try:
            from docx import Document
        except ImportError as exc:
            raise ImportError("Install python-docx to parse .docx files.") from exc

        docx_document = Document(document)
        return "\n\n".join(paragraph.text for paragraph in docx_document.paragraphs)
